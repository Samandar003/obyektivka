from docx import Document
from docx.shared import Inches
from .models import OutFile

def all_qarindoshlari(data:dict):
    key_column_1 = ['qarindoshligi_1', 'full_name_1', 'tugilgan_yili_joyi_1',
    'ish_joyi_1', 'turar_joyi_1']
    key_column_2 = ['qarindoshligi_2', 'full_name_2', 'tugilgan_yili_joyi_2',
    'ish_joyi_2', 'turar_joyi_2']
    key_column_3 = ['qarindoshligi_3', 'full_name_3', 'tugilgan_yili_joyi_3',
    'ish_joyi_3', 'turar_joyi_3']
    key_column_4 = ['qarindoshligi_4', 'full_name_4', 'tugilgan_yili_joyi_4',
    'ish_joyi_4', 'turar_joyi_4']
    key_column_5 = ['qarindoshligi_5', 'full_name_5', 'tugilgan_yili_joyi_5',
    'ish_joyi_5', 'turar_joyi_5']
    key_column_6 = ['qarindoshligi_6', 'full_name_6', 'tugilgan_yili_joyi_6',
    'ish_joyi_6', 'turar_joyi_6']
    new1 = [m for n, m in data.items() if n in key_column_1]
    new2 = [m for n, m in data.items() if n in key_column_2]
    new3 = [m for n, m in data.items() if n in key_column_3]
    new4 = [m for n, m in data.items() if n in key_column_4]
    new5 = [m for n, m in data.items() if n in key_column_5]
    new6 = [m for n, m in data.items() if n in key_column_6]
    result = []
    result.append(new1)
    result.append(new2)
    result.append(new3)
    result.append(new4)
    result.append(new5)
    result.append(new6)
    return result
            
def first_table__first_col_key(data:dict):
    default_keys = ['date_birth', 'nationality', 'place_birth', 'partiyaviyligi',
        'malumoti', 'malumoti_boyicha_mutaxasisligi', 'ilmiy_darajasi',
        'qaysi_chet_tillarini_biladi', 'tamomlagan_maktab', 'ilmiy_unvoni',
        'davlat_mukofotlari_tqadirlangan', 'mehnat_faoliyati', 'saylov_azosi']
    new = {}
    for n, m in data.items():
        if n in default_keys:
            new[n] = m
    return new


def intro_oby(full_name, picture, dic, qarindoshlari_keys):
    doc = Document()
    p = doc.add_heading('\t\tMalumotnoma\t\t\t', 1)
    r = p.add_run()
    r.add_picture(picture, width=Inches(1.25), height=Inches(1.25))
    doc.add_paragraph(f"\n\t\t{full_name.title()}\n")
    nomi = f"media/files/{full_name}_obyektivka.docx"
    
    doc_keys = ["Tug'ilgan yili", 'Millati', "Tug'ilgan joyi",
        "Partiyaviyligi", "Ma'lumoti", "Malumoti bo'yicha mutaxassisligi",
        "Ilmiy darajasi", "Qaysi chet tillari", "Tamomlagan maktab", 
        "Ilmiy unvoni", "Davlat Mukofotlari", "Mehnat faoliyati", "Saylov azosi"]
    
    dalniy_data = [x for x in dic.values()]
    table1 = doc.add_table(rows=13, cols=2)
    
    m = 0
    for cell in table1.columns[0].cells:
        cell.text = doc_keys[m]
        m += 1
    
    n = 0
    for cell in table1.columns[1].cells:
        cell.text = dalniy_data[n]
        n += 1
    
    doc.add_heading(f"\n\t{full_name.title()}ning eng yqain qarindoshlari hqaida malumot", 3)
    table2= doc.add_table(rows=7, cols=5)

    # Qarindoshlari boshlandi
    row = table2.rows[0].cells
    qarindoshlar = ['\nQarindoshligi', '\nFamiliyasi', '\nTugulgan yili va joyi', '\nIsh joyi', '\nTurar joyi']
    for n in list(range(len(qarindoshlar))):
        row[n].text = qarindoshlar[n]
    
    a = qarindoshlari_keys[0]
    b = qarindoshlari_keys[1]
    c = qarindoshlari_keys[2]
    d = qarindoshlari_keys[3]
    e = qarindoshlari_keys[4]
    f = qarindoshlari_keys[5]
 
    n = 0
    for cell in table2.rows[1].cells:
        cell.text = a[n]
        n += 1
    
    m = 0
    for cell in table2.rows[2].cells:
        cell.text = b[m]
        m += 1
        
    x = 0
    for cell in table2.rows[3].cells:
        cell.text = c[x]
        x += 1
    
    l = 0
    for cell in table2.rows[4].cells:
        cell.text = d[l]
        l += 1
    
    z = 0
    for cell in table2.rows[5].cells:
        cell.text = e[z]
        z += 1
    
    w = 0
    for cell in table2.rows[6].cells:
        cell.text = f[w]
        w += 1
    
    doc.save(f"media/files/{full_name}_obyektivka.docx")
    outfile = OutFile.objects.create(file=f"files/{full_name}_obyektivka.docx")
    outfile.save()
    return {"file":outfile.file}
    





